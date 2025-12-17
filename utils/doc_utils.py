from docx import Document
from docx.shared import Inches, Pt
import pandas as pd
import io
import matplotlib.pyplot as plt
from pandas.plotting import table

def clean_and_trim_df(df):
    """
    Cleans the DataFrame to remove empty rows/cols and whitespace.
    Returns a dataframe containing ONLY the used range.
    """
    # 1. Replace strings that are just whitespace with NaN
    df = df.replace(r'^\s*$', pd.NA, regex=True)
    
    # 2. Drop rows that are ENTIRELY empty (NaN)
    df = df.dropna(how='all', axis=0)
    
    # 3. Drop columns that are ENTIRELY empty
    df = df.dropna(how='all', axis=1)
    
    # 4. Fill remaining NaNs with empty strings for clean display
    df = df.fillna('')
    
    return df

def extract_valuation_table_data(excel_file):
    """
    Analyzes an Excel file (DCF/NAV), extracts the table data as a DataFrame.
    IMPORTANT: Reads with header=None to preserve exact grid layout.
    """
    try:
        xls = pd.ExcelFile(excel_file)
        sheet_names = xls.sheet_names
        target_df = None
        
        # --- Sheet Selection Logic ---
        
        # 1. Check for DCF -> Financials
        dcf_sheets = [s for s in sheet_names if 'DCF' in s.upper()]
        if dcf_sheets:
            if 'Financials' in sheet_names:
                # Read without header to keep the exact look (title, spacers, etc)
                target_df = pd.read_excel(xls, 'Financials', header=None)
            
        # 2. Check for NAV (Fallback)
        if target_df is None:
            nav_sheets = [s for s in sheet_names if 'NAV' in s.upper()]
            if nav_sheets:
                target_df = pd.read_excel(xls, nav_sheets[0], header=None)

        if target_df is None:
            return None

        # --- Data Cleaning ---
        target_df = clean_and_trim_df(target_df)
        
        if target_df.empty:
            return None

        return target_df

    except Exception as e:
        print(f"Error extracting table data: {e}")
        return None

def generate_financial_table_image(df):
    """
    Converts a dataframe into a high-quality 'Financial Report' style image buffer.
    Matches the 'Snapshot' look: Bold headers, clean lines, whitespace.
    """
    # Clean data: Replace NaN with empty strings
    df = df.fillna('')

    # Dynamic sizing: width based on cols, height based on rows
    # We estimate dimensions to ensure text doesn't wrap awkwardly
    w = max(len(df.columns) * 2.0, 8) 
    h = max((len(df) + 1) * 0.4, 3)
    
    fig, ax = plt.subplots(figsize=(w, h))
    ax.axis('off') # Hide the graph axes

    # Create the table
    # loc='center' centers it in the figure
    tbl = table(ax, df, loc='center', cellLoc='center', colWidths=[1.0/len(df.columns)] * len(df.columns))

    # Apply "Financial Report" Styling
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(11)
    tbl.scale(1.2, 1.5) # Scale width, height

    # Iterate over cells to apply specific styles
    for (row, col), cell in tbl.get_celld().items():
        cell.set_edgecolor('white') # Default to invisible borders
        cell.set_linewidth(0)
        
        # --- HEADERS (Row 0) ---
        if row == 0:
            cell.set_text_props(weight='bold', color='black')
            cell.set_facecolor('white')
            # Thick bottom border for header
            cell.set_edgecolor('black') 
            cell.set_linewidth(1.5)
            cell.visible_edges = "B" 
        
        # --- DATA ROWS ---
        else:
            cell.set_facecolor('white')
            cell.set_text_props(color='#333333')
            
            # Align logic: Attempt to right-align numbers
            val = df.iloc[row-1, col]
            
            # Simple heuristic for numbers vs text
            is_number = False
            try:
                if isinstance(val, (int, float)):
                    is_number = True
                elif isinstance(val, str):
                    # Clean chars to check if it's a number like "$50,000" or "(500)"
                    clean_val = val.replace('$', '').replace('%', '').replace(',', '').replace('(', '').replace(')', '').strip()
                    if clean_val and clean_val.replace('.', '', 1).isdigit():
                        is_number = True
            except:
                is_number = False

            if is_number:
                 cell.set_text_props(ha='right')
            else:
                 cell.set_text_props(ha='left')

            # Thin bottom border for rows
            cell.set_edgecolor('#e0e0e0')
            cell.set_linewidth(0.5)
            cell.visible_edges = "B"

    # Save to Buffer
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=300, pad_inches=0.1)
    buf.seek(0)
    plt.close(fig) # Close plot to free memory
    
    return buf

def process_word_template(template_path, context, table_data=None, provided_image_stream=None):
    """
    Opens a Word template, replaces placeholders.
    Inserts a table image from EITHER:
    1. 'provided_image_stream' (Direct user upload)
    2. 'table_data' (Auto-generated from DataFrame)
    """
    doc = Document(template_path)
    
    # Determine which image source to use
    final_image_stream = None
    
    if provided_image_stream is not None:
        # Use the manually uploaded image
        final_image_stream = provided_image_stream
    elif table_data is not None:
        # Generate the image from the dataframe
        final_image_stream = generate_financial_table_image(table_data)

    # 1. Iterate through all paragraphs
    paragraphs = list(doc.paragraphs)
    
    for p in paragraphs:
        # --- Table Image Insertion Logic ---
        if '<<valuation.jpg>>' in p.text and final_image_stream is not None:
            p.text = "" # Clear the placeholder text
            run = p.add_run()
            # Add picture with a set width (e.g., 6 inches to fit standard page margins)
            # You can adjust width=Inches(6.0) if you want it larger/smaller
            run.add_picture(final_image_stream, width=Inches(6.0))
            continue

        # --- Standard Text Replacement ---
        for key, value in context.items():
            if key in p.text:
                p.text = p.text.replace(key, str(value))

    # 2. Iterate through existing tables to replace text inside them
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                     for key, value in context.items():
                        if key in p.text:
                            p.text = p.text.replace(key, str(value))
    
    return doc